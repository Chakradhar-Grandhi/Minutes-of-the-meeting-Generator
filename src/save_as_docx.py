from docx import Document

def save_as_docx(transcript, summary, sentiment_results, file_name='Meeting_Minutes.docx'):
    print("Saving Doc....")
    doc = Document()
    doc.add_heading('Meeting Minutes', 0)

    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)

    doc.add_heading('Transcript', level=1)
    doc.add_paragraph(transcript)

    doc.add_heading('Sentiment Analysis', level=1)
    for result in sentiment_results:
        doc.add_paragraph(f"Text: {result['text']}\nSentiment: {result['label']} ({result['score']:.2f})\n")

    doc.save(file_name)
    print("Doc Saved....")
